from typing import List, Dict, Tuple, Optional
from pathlib import Path
from docx import Document
from docx.shared import Pt
import csv
from langchain_text_splitters import RecursiveCharacterTextSplitter


class MarkdownProcessorLocal:
    def __init__(self,
                 chunk_size: int = 300,
                 default_output_dir: Path | str | None = None):
        self.chunk_size = chunk_size
        self.current_headers = {1: "", 2: "", 3: "", 4: "", 5: "", 6: ""}
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=0,
            separators=["\n\n", "\n", "。", "；", "，", " ", ""]
        )

        # 设置默认输出目录
        if default_output_dir is None:
            self.default_output_dir = Path.cwd() / 'output'
        else:
            self.default_output_dir = Path(default_output_dir)

    def get_header_level(self, line: str) -> int:
        if not line.strip().startswith('#'):
            return 0
        level = 0
        for char in line.strip():
            if char == '#':
                level += 1
            else:
                break
        return level

    def update_headers(self, line: str, level: int):
        header_text = line.lstrip('#').strip()
        self.current_headers[level] = header_text
        for i in range(level + 1, 7):
            self.current_headers[i] = ""

    def collect_current_headers(self) -> str:
        headers = []
        for level in range(1, 7):
            if self.current_headers[level]:
                headers.append("#" * level + " " + self.current_headers[level])
        return "\n".join(headers) + "\n\n" if headers else ""

    def is_table_separator(self, line: str) -> bool:
        return bool(line.strip().startswith('|') and '----' in line)

    def is_table_row(self, line: str) -> bool:
        return bool(line.strip().startswith('|') and line.strip().endswith('|'))

    def split_table_row(self, row: str) -> List[str]:
        return [cell.strip() for cell in row.strip('|').split('|')]

    def join_table_row(self, cells: List[str]) -> str:
        return '| ' + ' | '.join(cells) + ' |'

    def get_first_column(self, row: str) -> str:
        cells = self.split_table_row(row)
        return cells[0] if cells else ""

    def split_table(self, table_lines: List[str]) -> List[List[str]]:
        if len(table_lines) < 3:
            return [table_lines]

        header_row = table_lines[0]
        separator_row = table_lines[1]
        content_rows = table_lines[2:]

        if len('\n'.join(table_lines)) <= self.chunk_size:
            return [table_lines]

        rows_per_chunk = max(2, (self.chunk_size - len(header_row) - len(separator_row)) //
                             (len(content_rows[0]) if content_rows else 1))

        split_tables = []
        for i in range(0, len(content_rows), rows_per_chunk):
            chunk_rows = content_rows[i:i + rows_per_chunk]
            split_tables.append([
                header_row,
                separator_row,
                *chunk_rows
            ])

        return split_tables

    def extract_table(self, lines: List[str], start_idx: int) -> Tuple[List[str], int]:
        table_lines = []
        i = start_idx

        while i >= 0 and (self.is_table_row(lines[i]) or self.is_table_separator(lines[i])):
            table_lines.insert(0, lines[i])
            i -= 1

        i = start_idx + 1
        while i < len(lines) and (self.is_table_row(lines[i]) or self.is_table_separator(lines[i])):
            table_lines.append(lines[i])
            i += 1

        return table_lines, i - 1

    def process_markdown(self, markdown_text: str):
        lines = markdown_text.split('\n')
        content_buffer = []
        headers = ""
        i = 0

        while i < len(lines):
            line = lines[i].rstrip()

            if not line:
                content_buffer.append(line)
                i += 1
                continue

            level = self.get_header_level(line)
            if level > 0:
                if content_buffer:
                    text = "\n".join(content_buffer)
                    if text.strip():
                        chunks = self.text_splitter.split_text(text.strip())
                        for chunk in chunks:
                            yield headers + chunk + "\n"
                    content_buffer = []

                self.update_headers(line, level)
                headers = self.collect_current_headers()
            elif self.is_table_separator(line) or self.is_table_row(line):
                table_lines, end_idx = self.extract_table(lines, i)

                if content_buffer:
                    text = "\n".join(content_buffer)
                    if text.strip():
                        chunks = self.text_splitter.split_text(text.strip())
                        for chunk in chunks:
                            yield headers + chunk + "\n"
                    content_buffer = []

                split_tables = self.split_table(table_lines)
                for table in split_tables:
                    yield headers + "\n".join(table) + "\n\n"

                i = end_idx + 1
                continue
            else:
                content_buffer.append(line)

            i += 1

        if content_buffer:
            text = "\n".join(content_buffer)
            if text.strip():
                chunks = self.text_splitter.split_text(text.strip())
                for chunk in chunks:
                    yield headers + chunk + "\n"

    def create_output_files(
            self,
            input_path: str | Path,
            output_dir: Optional[str | Path] = None,
            create_md: bool = True,
            create_docx: bool = True,
            create_csv: bool = True
    ) -> List[str]:
        """处理Markdown文件并保存为多种格式

        Args:
            input_path: 输入文件的完整路径或相对路径
            output_dir: 输出目录的路径，如果为None则使用默认的output目录
            create_md: 是否创建markdown输出
            create_docx: 是否创建word文档输出
            create_csv: 是否创建csv输出

        Returns:
            List[str]: 处理后的文本块列表

        Raises:
            FileNotFoundError: 当输入文件不存在时
            ValueError: 当输入文件不是markdown文件时
            Exception: 其他处理过程中的错误
        """
        try:
            # 处理输入路径
            input_path = Path(input_path)
            if not input_path.exists():
                raise FileNotFoundError(f"找不到输入文件: {input_path}")

            if input_path.suffix.lower() not in ['.md', '.markdown']:
                raise ValueError(f"不支持的文件格式: {input_path.suffix}")

            # 处理输出目录
            if output_dir is None:
                output_dir = self.default_output_dir / input_path.stem
            else:
                output_dir = Path(output_dir)

            output_dir.mkdir(parents=True, exist_ok=True)

            # 读取输入文件
            with open(input_path, 'r', encoding='utf-8') as f:
                markdown_text = f.read()

            # 处理文本
            results = list(self.process_markdown(markdown_text))

            # 定义分隔符
            separator = "=" * 40

            # 1. 保存 Markdown 文件
            if create_md:
                md_output = output_dir / f"{input_path.stem}.md"
                with open(md_output, 'w', encoding='utf-8') as f:
                    for chunk in results:
                        f.write(f"\n{separator}\n\n")
                        f.write(chunk)
                print(f"已生成 Markdown 文件：{md_output}")

            # 2. 创建 Word 文档
            if create_docx:
                docx_output = output_dir / f"{input_path.stem}.docx"
                doc = Document()
                for chunk in results:
                    separator_paragraph = doc.add_paragraph()
                    separator_paragraph.add_run('=' * 40)
                    content_paragraph = doc.add_paragraph()
                    content_paragraph.add_run(chunk)
                    for run in separator_paragraph.runs + content_paragraph.runs:
                        run.font.size = Pt(11)
                        run.font.name = 'Arial'
                doc.save(docx_output)
                print(f"已生成 Word 文件：{docx_output}")

            # 3. 保存 CSV 文件
            if create_csv:
                csv_output = output_dir / f"{input_path.stem}.csv"
                with open(csv_output, 'w', encoding='utf-8', newline='') as f:
                    writer = csv.writer(f)
                    writer.writerow(['Content'])
                    for chunk in results:
                        writer.writerow([chunk.strip()])
                print(f"已生成 CSV 文件：{csv_output}")

            print(f"\n处理完成！共生成 {len(results)} 个文本块")
            return results

        except Exception as e:
            print(f"处理过程中发生错误：{str(e)}")
            raise